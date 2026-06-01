from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
IGNORED_DIR_NAMES = {".venv", ".git", "__pycache__", "node_modules"}


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str


def list_supported_files(base_path: Path, recursive: bool = True) -> List[Path]:
    base = Path(base_path).expanduser().resolve()
    if not base.exists():
        raise FileNotFoundError(f"Caminho nao encontrado: {base}")
    if base.is_file():
        return [base] if base.suffix.lower() in SUPPORTED_EXTENSIONS else []
    if not base.is_dir():
        raise NotADirectoryError(f"Caminho nao e uma pasta: {base}")

    candidates = base.rglob("*") if recursive else base.glob("*")
    files: List[Path] = []
    for candidate in candidates:
        if not candidate.is_file():
            continue
        if candidate.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if any(part in IGNORED_DIR_NAMES for part in candidate.parts):
            continue
        files.append(candidate)
    return sorted(files)


def extract_text(path: Path) -> str:
    file_path = Path(path).expanduser().resolve()
    extension = file_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Extensao nao suportada: {extension}")

    if extension == ".pdf":
        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    if extension == ".docx":
        document = DocxDocument(str(file_path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()

    return file_path.read_text(encoding="utf-8", errors="ignore").strip()


def load_documents(base_path: Path, recursive: bool = True) -> List[LoadedDocument]:
    documents: List[LoadedDocument] = []
    for file_path in list_supported_files(base_path, recursive=recursive):
        text = extract_text(file_path)
        if text:
            documents.append(LoadedDocument(path=file_path, text=text))
    return documents
